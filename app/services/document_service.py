"""
Servicio de Documentos.

Orquesta la generación de documentos desde plantillas .docx y
registra cada generación en el historial.
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx2pdf import convert

from app.core.entities import DocumentoGeneradoDTO, VehiculoDTO
from app.infrastructure.repositories.documento_repository import DocumentoRepository
from app.infrastructure.repositories.vehiculo_repository import VehiculoRepository
from app.utils.path_manager import PathManager


def _reemplazar_texto_en_parrafo(parrafo, datos: Dict[str, str]) -> None:
    """
    Infraestructura base del motor run-aware (RPSR).

    Sprint 5B.1: solo construye texto lógico y mapa rico por carácter.
    No detecta placeholders ni modifica runs.
    """

    # Fase 0: salida temprana para evitar procesamiento innecesario.
    runs = list(parrafo.runs)
    if not runs:
        return

    texto_logico = "".join((run.text or "") for run in runs)
    if "{" not in texto_logico:
        return

    # Fase 1 + Fase 2: mapa rico por carácter y texto lógico concatenado.
    # Cada entrada conserva trazabilidad exacta entre índice global y run original.
    mapa_caracteres = []
    indice_global = 0
    for indice_run, run in enumerate(runs):
        texto_run = run.text or ""
        for posicion_run, caracter in enumerate(texto_run):
            mapa_caracteres.append(
                {
                    "char": caracter,
                    "global_index": indice_global,
                    "run_index": indice_run,
                    "run_char_index": posicion_run,
                    "run_ref": run,
                }
            )
            indice_global += 1

    # Validación interna de infraestructura base (sin efectos en documento).
    if len(mapa_caracteres) != len(texto_logico):
        return

    # Fase 3: escáner secuencial de placeholders sobre texto lógico.
    # Solo detecta spans válidos; no modifica texto ni runs.
    def _es_nombre_variable_valido(nombre: str) -> bool:
        if not nombre:
            return False
        return all(ch.isalnum() or ch == "_" for ch in nombre)

    spans = []
    longitud = len(texto_logico)
    i = 0

    while i < longitud:
        if texto_logico[i] != "{":
            i += 1
            continue

        # Caso doble llave: {{variable}}
        if i + 1 < longitud and texto_logico[i + 1] == "{":
            inicio = i
            j = i + 2
            encontrado = False

            while j + 1 < longitud:
                if texto_logico[j] == "}" and texto_logico[j + 1] == "}":
                    nombre_variable = texto_logico[i + 2 : j]
                    if _es_nombre_variable_valido(nombre_variable):
                        placeholder = texto_logico[inicio : j + 2]
                        spans.append(
                            {
                                "start_index": inicio,
                                "end_index": j + 1,
                                "placeholder_text": placeholder,
                                "variable_name": nombre_variable,
                                "placeholder_length": len(placeholder),
                            }
                        )
                        i = j + 2
                        encontrado = True
                    else:
                        i = inicio + 2
                        encontrado = True
                    break
                j += 1

            if not encontrado:
                i = inicio + 2
            continue

        # Caso llave simple: {variable}
        inicio = i
        j = i + 1
        while j < longitud and texto_logico[j] != "}":
            if texto_logico[j] == "{":
                break
            j += 1

        if j < longitud and texto_logico[j] == "}":
            nombre_variable = texto_logico[i + 1 : j]
            if _es_nombre_variable_valido(nombre_variable):
                placeholder = texto_logico[inicio : j + 1]
                spans.append(
                    {
                        "start_index": inicio,
                        "end_index": j,
                        "placeholder_text": placeholder,
                        "variable_name": nombre_variable,
                        "placeholder_length": len(placeholder),
                    }
                )
                i = j + 1
                continue

        # Placeholder dañado o incompleto: se ignora.
        i = inicio + 1

    # Fase 4: resolver spans lógicos hacia runs físicos y offsets exactos.
    # Solo calcula metadatos para fases posteriores; no modifica el documento.
    placeholders_resueltos = []
    for span in spans:
        inicio = span["start_index"]
        fin = span["end_index"]

        if inicio < 0 or fin >= len(mapa_caracteres) or inicio > fin:
            continue

        entrada_inicio = mapa_caracteres[inicio]
        entrada_fin = mapa_caracteres[fin]

        start_run_index = entrada_inicio["run_index"]
        end_run_index = entrada_fin["run_index"]
        start_offset = entrada_inicio["run_char_index"]
        end_offset = entrada_fin["run_char_index"]

        tramo = mapa_caracteres[inicio : fin + 1]
        conteo_por_run = {}
        for entrada in tramo:
            indice_run = entrada["run_index"]
            conteo_por_run[indice_run] = conteo_por_run.get(indice_run, 0) + 1

        affected_run_indices = sorted(conteo_por_run.keys())
        affected_run_count = len(affected_run_indices)

        if affected_run_count == 1:
            base_run_index = affected_run_indices[0]
        elif affected_run_count == 2:
            base_run_index = affected_run_indices[0]
        else:
            base_run_index = max(
                affected_run_indices,
                key=lambda indice: (conteo_por_run[indice], -indice),
            )

        placeholders_resueltos.append(
            {
                "placeholder_text": span["placeholder_text"],
                "variable_name": span["variable_name"],
                "start_index": inicio,
                "end_index": fin,
                "start_run_index": start_run_index,
                "end_run_index": end_run_index,
                "start_offset": start_offset,
                "end_offset": end_offset,
                "affected_run_count": affected_run_count,
                "base_run_index": base_run_index,
            }
        )

    # Fase 5: aplicar reemplazo físico sobre runs sin reconstruir el párrafo.
    # Procesa de derecha a izquierda para mantener válidos los offsets previos.
    for placeholder in sorted(
        placeholders_resueltos,
        key=lambda item: item["start_index"],
        reverse=True,
    ):
        nombre_variable = placeholder["variable_name"]
        if nombre_variable not in datos:
            continue

        valor_reemplazo = str(datos[nombre_variable])
        start_run_index = placeholder["start_run_index"]
        end_run_index = placeholder["end_run_index"]
        start_offset = placeholder["start_offset"]
        end_offset = placeholder["end_offset"]
        base_run_index = placeholder["base_run_index"]

        if start_run_index == end_run_index:
            run = runs[start_run_index]
            texto_run = run.text or ""
            prefijo = texto_run[:start_offset]
            sufijo = texto_run[end_offset + 1 :]
            run.text = f"{prefijo}{valor_reemplazo}{sufijo}"
            continue

        for indice_run in range(start_run_index, end_run_index + 1):
            run = runs[indice_run]
            texto_run = run.text or ""

            if indice_run == start_run_index:
                nuevo_texto = texto_run[:start_offset]
                if indice_run == base_run_index:
                    nuevo_texto += valor_reemplazo
            elif indice_run == end_run_index:
                nuevo_texto = texto_run[end_offset + 1 :]
                if indice_run == base_run_index:
                    nuevo_texto = f"{valor_reemplazo}{nuevo_texto}"
            elif indice_run == base_run_index:
                nuevo_texto = valor_reemplazo
            else:
                nuevo_texto = ""

            run.text = nuevo_texto


def _reemplazar_texto_en_tabla(tabla, datos: Dict[str, str]) -> None:
    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                _reemplazar_texto_en_parrafo(parrafo, datos)
            for tabla_anidada in celda.tables:
                _reemplazar_texto_en_tabla(tabla_anidada, datos)


def _reemplazar_texto_en_documento(documento: Document, datos: Dict[str, str]) -> None:
    for parrafo in documento.paragraphs:
        _reemplazar_texto_en_parrafo(parrafo, datos)
    for tabla in documento.tables:
        _reemplazar_texto_en_tabla(tabla, datos)
    for section in documento.sections:
        for parrafo in section.header.paragraphs:
            _reemplazar_texto_en_parrafo(parrafo, datos)
        for tabla in section.header.tables:
            _reemplazar_texto_en_tabla(tabla, datos)
        for parrafo in section.footer.paragraphs:
            _reemplazar_texto_en_parrafo(parrafo, datos)
        for tabla in section.footer.tables:
            _reemplazar_texto_en_tabla(tabla, datos)


def _detectar_marcadores_residuales_en_texto(texto: str) -> List[str]:
    hallazgos = []
    longitud = len(texto)
    mascara = [False] * longitud
    i = 0

    while i < longitud:
        if texto[i] != "{":
            i += 1
            continue

        if i + 1 < longitud and texto[i + 1] == "{":
            inicio = i
            j = i + 2
            encontrado = False

            while j + 1 < longitud:
                if texto[j] == "}" and texto[j + 1] == "}":
                    nombre = texto[i + 2 : j]
                    if nombre and all(ch.isalnum() or ch == "_" for ch in nombre):
                        placeholder = texto[inicio : j + 2]
                        hallazgos.append(f"placeholder residual: {placeholder}")
                        for indice in range(inicio, j + 2):
                            mascara[indice] = True
                        i = j + 2
                        encontrado = True
                    else:
                        i = inicio + 2
                        encontrado = True
                    break
                j += 1

            if not encontrado:
                i = inicio + 2
            continue

        inicio = i
        j = i + 1
        while j < longitud and texto[j] != "}":
            if texto[j] == "{":
                break
            j += 1

        if j < longitud and texto[j] == "}":
            nombre = texto[i + 1 : j]
            if nombre and all(ch.isalnum() or ch == "_" for ch in nombre):
                placeholder = texto[inicio : j + 1]
                hallazgos.append(f"placeholder residual: {placeholder}")
                for indice in range(inicio, j + 1):
                    mascara[indice] = True
                i = j + 1
                continue

        i = inicio + 1

    residuo_llaves = "".join(
        caracter
        for indice, caracter in enumerate(texto)
        if not mascara[indice] and caracter in "{}"
    )
    if residuo_llaves:
        hallazgos.append(f"llaves residuales: {residuo_llaves}")

    return hallazgos


def _registrar_hallazgos_en_parrafos(parrafos, ubicacion_base: str, hallazgos: List[str]) -> None:
    for indice, parrafo in enumerate(parrafos, start=1):
        texto = parrafo.text or ""
        if "{" not in texto and "}" not in texto:
            continue

        hallazgos_texto = _detectar_marcadores_residuales_en_texto(texto)
        for hallazgo in hallazgos_texto:
            hallazgos.append(f"{ubicacion_base} párrafo {indice}: {hallazgo}")


def _registrar_hallazgos_en_tabla(tabla, ubicacion_base: str, hallazgos: List[str]) -> None:
    for indice_fila, fila in enumerate(tabla.rows, start=1):
        for indice_celda, celda in enumerate(fila.cells, start=1):
            ubicacion_celda = f"{ubicacion_base} fila {indice_fila} celda {indice_celda}"
            _registrar_hallazgos_en_parrafos(celda.paragraphs, ubicacion_celda, hallazgos)
            for indice_tabla, tabla_anidada in enumerate(celda.tables, start=1):
                ubicacion_tabla = f"{ubicacion_celda} tabla anidada {indice_tabla}"
                _registrar_hallazgos_en_tabla(tabla_anidada, ubicacion_tabla, hallazgos)


def _validar_documento_generado(documento: Document) -> None:
    hallazgos = []

    _registrar_hallazgos_en_parrafos(documento.paragraphs, "body", hallazgos)
    for indice_tabla, tabla in enumerate(documento.tables, start=1):
        _registrar_hallazgos_en_tabla(tabla, f"body tabla {indice_tabla}", hallazgos)

    for indice_seccion, section in enumerate(documento.sections, start=1):
        _registrar_hallazgos_en_parrafos(
            section.header.paragraphs,
            f"header sección {indice_seccion}",
            hallazgos,
        )
        for indice_tabla, tabla in enumerate(section.header.tables, start=1):
            _registrar_hallazgos_en_tabla(
                tabla,
                f"header sección {indice_seccion} tabla {indice_tabla}",
                hallazgos,
            )

        _registrar_hallazgos_en_parrafos(
            section.footer.paragraphs,
            f"footer sección {indice_seccion}",
            hallazgos,
        )
        for indice_tabla, tabla in enumerate(section.footer.tables, start=1):
            _registrar_hallazgos_en_tabla(
                tabla,
                f"footer sección {indice_seccion} tabla {indice_tabla}",
                hallazgos,
            )

    if hallazgos:
        detalle = " | ".join(hallazgos)
        raise ValueError(
            f"Documento incompleto. Hallazgos: {len(hallazgos)}. {detalle}"
        )


def generar_documento(ruta_plantilla: Path, ruta_salida: Path, datos: Dict[str, str]) -> Path:
    if not ruta_plantilla.exists():
        raise FileNotFoundError(f"Plantilla no encontrada: {ruta_plantilla}")

    documento = Document(ruta_plantilla)
    _reemplazar_texto_en_documento(documento, datos)
    _validar_documento_generado(documento)
    documento.save(ruta_salida)
    return ruta_salida


class DocumentoService:
    """Coordina la generación de documentos y el historial."""

    def __init__(
        self,
        vehiculo_repo: VehiculoRepository = None,
        documento_repo: DocumentoRepository = None,
    ):
        self.vehiculo_repo = vehiculo_repo or VehiculoRepository()
        self.documento_repo = documento_repo or DocumentoRepository()
        self.config = self._cargar_configuracion()

    def _cargar_configuracion(self) -> dict:
        """Carga la configuración de tipos de documentos desde el JSON."""
        ruta_config = PathManager.get_documents_config_path()
        print("\n" + "=" * 70)
        print("DEBUG _cargar_configuracion()")
        print("Ruta absoluta JSON:", ruta_config.resolve())
        with open(ruta_config, "r", encoding="utf-8") as f:
            config = json.load(f)

        print("Tipos de documento cargados:")
        for tipo_documento in config.keys():
            print("-", tipo_documento)

        tipo_objetivo = "Informe Accidente Daños Materiales"
        campos_objetivo = config.get(tipo_objetivo, {}).get("campos_editables", [])
        print(f"Campos editables de '{tipo_objetivo}':")
        print(campos_objetivo)
        print("=" * 70 + "\n")

        return config

    def obtener_tipos_documento(self) -> List[str]:
        """Devuelve la lista de tipos de documento disponibles."""
        return list(self.config.keys())

    def obtener_configuracion_documento(self, tipo: str) -> Optional[dict]:
        """Devuelve la configuración para un tipo de documento específico."""
        return self.config.get(tipo)

    def cargar_datos_vehiculo(self, vehiculo_id: int) -> Optional[VehiculoDTO]:
        """Obtiene los datos completos de un vehículo por su ID."""
        from app.services.vehiculo_mysql_service import VehiculoMySQLService

        return VehiculoMySQLService().buscar_por_id(vehiculo_id)

    # =================== MÉTODOS AUXILIARES PARA PREPARAR DATOS ===================

    def _obtener_fecha_actual(self) -> Tuple[str, str]:
        """Devuelve la fecha actual en formato corto y largo."""
        ahora = datetime.now()
        meses_es = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        hoy = ahora.strftime("%d/%m/%Y")
        fecha_larga = f"{ahora.strftime('%d')} de {meses_es[ahora.month - 1]} de {ahora.strftime('%Y')}"
        return hoy, fecha_larga

    def _datos_sistema(self) -> Dict[str, str]:
        """
        Genera las variables automáticas del sistema a partir de la fecha y hora actuales.

        Retorna un diccionario con:
            - dia         : día del mes (número)
            - mes         : nombre completo del mes en español
            - anio        : año con cuatro dígitos
            - fecha_corta : DD/MM/AAAA
            - fecha_larga : "día de mes de año"

        Esta función centraliza la creación de variables de sistema y puede extenderse
        en el futuro con nuevas variables (ej. hora, minuto, etc.) sin afectar al resto.
        """
        ahora = datetime.now()
        meses_es = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        return {
            "dia": ahora.strftime("%d"),
            "mes": meses_es[ahora.month - 1],
            "anio": ahora.strftime("%Y"),
            "fecha_corta": ahora.strftime("%d/%m/%Y"),
            "fecha_larga": f"{ahora.strftime('%d')} de {meses_es[ahora.month - 1]} de {ahora.strftime('%Y')}",
        }

    def _datos_vehiculo(self, vehiculo: VehiculoDTO) -> Dict[str, str]:
        """Extrae los campos del vehículo (excepto propietario y conductor)."""
        # NOTA: numero_motor es solo el motor real; nunca se usa numero_interno como fallback.
        return {
            "placa": vehiculo.placa or "",
            "numero_interno": str(vehiculo.numero_interno) if vehiculo.numero_interno else "",
            "marca": vehiculo.marca or "",
            "modelo": vehiculo.modelo or "",
            "clase": vehiculo.clase or "",
            "motor": vehiculo.motor or "",
            "chasis": vehiculo.chasis or "",
            "serie": vehiculo.serie or "",
            "vin": vehiculo.vin or "",
            "fecha_matricula": vehiculo.fecha_matricula or "",
            "fecha_afiliacion": vehiculo.fecha_afiliacion or "",
            "capacidad": vehiculo.capacidad or "",
            "tipo": vehiculo.tipo or "",
            "combustible": vehiculo.combustible or "",
            "modalidad": vehiculo.modalidad or "",
            "ruta": vehiculo.ruta or "",
            "color": vehiculo.color or "",
            "carroceria": vehiculo.carroceria or "",
            "servicio": vehiculo.servicio or "",
            "numero_motor": vehiculo.motor or "",  # alias heredado, se retirará en futura versión
        }

    def _datos_propietario(self, vehiculo: VehiculoDTO) -> Dict[str, str]:
        """Extrae los campos del propietario."""
        # documento es un alias heredado para mantener compatibilidad con plantillas existentes.
        # documento_propietario será la variable oficial para todas las nuevas plantillas.
        return {
            "nombre_propietario": vehiculo.nombre_propietario or "",
            "documento": vehiculo.documento_propietario or "",  # alias heredado
            "documento_propietario": vehiculo.documento_propietario or "",
            "ciudad_expedicion": vehiculo.ciudad_expedicion or "",
            "telefono_propietario": vehiculo.telefono_propietario or "",
            "direccion_propietario": vehiculo.direccion_propietario or "",
        }

    def _datos_conductor(self, vehiculo: VehiculoDTO) -> Dict[str, str]:
        """Extrae los campos del conductor."""
        return {
            "nombre_conductor": vehiculo.nombre_conductor or "",
            "documento_conductor": vehiculo.documento_conductor or "",
            "celular_conductor": vehiculo.celular_conductor or "",
            "direccion_conductor": vehiculo.direccion_conductor or "",
            "correo_conductor": vehiculo.correo_conductor or "",
        }

    def _aplicar_editables(self, datos: Dict[str, str], datos_editables: Dict[str, str], hoy: str) -> None:
        """
        Procesa los campos editables.
        Si el campo es de tipo fecha y está vacío, se rellena con la fecha actual.
        """
        for campo, valor in datos_editables.items():
            if "fecha" in campo.lower() and not valor.strip():
                datos[campo] = hoy
            else:
                datos[campo] = valor.strip() if valor else ""

    def _aplicar_defaults_config(self, datos: Dict[str, str], tipo: str, hoy: str) -> None:
        """
        Rellena los campos faltantes que vienen definidos en la configuración
        para el tipo de documento.
        """
        config = self.config.get(tipo, {})
        for item in config.get("campos_editables", []):
            if len(item) < 3:
                continue
            label = item[0]          # Etiqueta visible (no se usa en lógica)
            clave = item[1]          # Nombre de la variable
            valor_defecto = item[2]  # Valor por defecto
            if clave not in datos:
                if "fecha" in clave.lower() and not valor_defecto:
                    datos[clave] = hoy
                else:
                    datos[clave] = valor_defecto if valor_defecto else ""

    def _formatear_fecha_larga(self, fecha: str) -> str:
        """
        Convierte una fecha en formato DD/MM/AAAA a formato largo en español:
        "15 de julio de 2026"

        Si la fecha no es válida o no tiene el formato esperado, devuelve el valor original.
        """
        meses = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        try:
            dia, mes, anio = map(int, fecha.split('/'))
            if 1 <= dia <= 31 and 1 <= mes <= 12 and 1000 <= anio <= 9999:
                return f"{dia} de {meses[mes-1]} de {anio}"
        except Exception:
            pass
        return fecha

    # =================== MÉTODO PRINCIPAL REFACTORIZADO ===================

    def preparar_datos_plantilla(
        self, vehiculo: VehiculoDTO, datos_editables: Dict[str, str], tipo: str
    ) -> Dict[str, str]:
        """
        Combina los datos automáticos del vehículo con los campos editables
        y prepara el diccionario final para la plantilla.
        """
        print("\n" + "=" * 70)
        print("DEBUG preparar_datos_plantilla() - diccionario recibido")
        print("Tipo documento:", tipo)
        print(datos_editables)
        print("Contiene fecha_reclamacion:", "fecha_reclamacion" in datos_editables)
        print("=" * 70)

        # Obtener variables del sistema (dia, mes, anio, fecha_corta, fecha_larga)
        datos_sistema = self._datos_sistema()
        hoy = datos_sistema["fecha_corta"]

        datos = {}

        # Bloque: datos del vehículo
        datos.update(self._datos_vehiculo(vehiculo))

        # Bloque: datos del propietario
        datos.update(self._datos_propietario(vehiculo))

        # Bloque: datos del conductor
        datos.update(self._datos_conductor(vehiculo))

        # Incorporar todas las variables de sistema
        datos.update(datos_sistema)

        # Aplicar campos editables (sobreescriben automáticos)
        self._aplicar_editables(datos, datos_editables, hoy)

        # Rellenar defaults desde configuración (para campos no presentes)
        self._aplicar_defaults_config(datos, tipo, hoy)

        # Convertir fechas en formato DD/MM/AAAA a formato largo (día de mes de año)
        for clave, valor in datos.items():
            # Evitar convertir variables del sistema que ya tienen formato largo o que deben mantenerse
            if clave in ("fecha_corta", "fecha_larga"):
                continue
            if "fecha" in clave.lower() and isinstance(valor, str) and len(valor) == 10:
                if valor[2] == '/' and valor[5] == '/':
                    datos[clave] = self._formatear_fecha_larga(valor)

        # ==========================================================
        # AUDITORÍA TEMPORAL – VALORES FINALES DE FECHAS
        # ==========================================================
        print("\n" + "=" * 70)
        print("AUDITORÍA FINAL DE DATOS ENVIADOS A LA PLANTILLA")
        print("=" * 70)
        for clave in sorted(datos.keys()):
            if "fecha" in clave.lower():
                print(f"{clave:25} -> {datos[clave]}")
        print("=" * 70 + "\n")
        # ==========================================================

        print("\n" + "=" * 70)
        print("DEBUG preparar_datos_plantilla() - diccionario final")
        print(datos)
        print("Contiene fecha_reclamacion:", "fecha_reclamacion" in datos)
        if "fecha_reclamacion" in datos:
            print("Valor fecha_reclamacion:", datos.get("fecha_reclamacion"))
        print("=" * 70 + "\n")

        # =================================================================
        # PREPARACIÓN PARA FUTURAS AMPLIACIONES
        # =================================================================
        # Este método es el punto central de integración de datos para
        # cualquier plantilla de documento.
        #
        # Actualmente combina los siguientes proveedores de datos:
        #   - Vehículo          → _datos_vehiculo()
        #   - Propietario       → _datos_propietario()
        #   - Conductor         → _datos_conductor()
        #   - Sistema           → _datos_sistema()
        #   - Editables         → _aplicar_editables()
        #   - Valores por defecto → _aplicar_defaults_config()
        #
        # Cada proveedor devuelve un diccionario independiente y se integra
        # exclusivamente mediante datos.update(), manteniendo el principio
        # de responsabilidad única (SRP).
        #
        # En futuros sprints se podrán añadir nuevos proveedores de datos,
        # por ejemplo:
        #   - _datos_accidente()
        #   - _datos_tercero()
        #   - _datos_lesionados()
        #   - _datos_occisos()
        #   - _datos_afectados()
        #
        # Cada nuevo proveedor deberá devolver un diccionario y su
        # integración se realizará mediante datos.update(), sin modificar
        # la estructura principal de este método.
        # =================================================================

        return datos

    # =================== GENERACIÓN Y REGISTRO ===================

    def generar_documento(
        self,
        tipo: str,
        vehiculo: VehiculoDTO,
        datos_editables: Dict[str, str],
        ruta_salida: str,
    ) -> Tuple[bool, str]:
        """
        Genera el documento Word y registra en el historial.

        Returns:
            Tupla (éxito, mensaje). Si éxito, el mensaje es la ruta del archivo generado.
        """
        config = self.config.get(tipo)
        if not config:
            return False, "Tipo de documento no válido."

        if not vehiculo:
            return False, "No se encontraron datos del vehículo."

        # Preparar datos combinados
        datos = self.preparar_datos_plantilla(vehiculo, datos_editables, tipo)

        # Ruta de la plantilla
        ruta_plantilla = PathManager.get_templates_dir() / config["plantilla"]

        try:
            salida_path = Path(ruta_salida)
            es_pdf = salida_path.suffix.lower() == ".pdf"

            if es_pdf:
                ruta_docx_temporal = salida_path.with_suffix(".docx")
                generar_documento(ruta_plantilla, ruta_docx_temporal, datos)
                convert(str(ruta_docx_temporal), str(salida_path.parent))

                ruta_pdf_final = salida_path.with_suffix(".pdf")
                if not ruta_pdf_final.exists():
                    raise FileNotFoundError(f"No se encontró el PDF generado: {ruta_pdf_final}")

                ruta_docx_temporal.unlink(missing_ok=True)

                registro = DocumentoGeneradoDTO(
                    tipo_documento=tipo,
                    placa_vehiculo=vehiculo.placa,
                    ruta_archivo=str(ruta_pdf_final),
                )
                self.documento_repo.insertar(registro)

                return True, str(ruta_pdf_final)

            # Generar documento físico
            ruta_generada = generar_documento(ruta_plantilla, salida_path, datos)

            # Registrar en historial
            registro = DocumentoGeneradoDTO(
                tipo_documento=tipo,
                placa_vehiculo=vehiculo.placa,
                ruta_archivo=str(ruta_generada),
            )
            self.documento_repo.insertar(registro)

            return True, str(ruta_generada)
        except FileNotFoundError as e:
            return False, f"No se encontró la plantilla: {e}"
        except Exception as e:
            return False, f"Error al generar el documento: {str(e)}"

    def obtener_historial(self) -> List[DocumentoGeneradoDTO]:
        """Devuelve el historial de documentos generados."""
        return self.documento_repo.obtener_todos()